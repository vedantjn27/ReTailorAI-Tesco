import os
import sys
import subprocess

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    from docx import Document
except ImportError:
    install('python-docx')
    from docx import Document

try:
    from reportlab.pdfgen import canvas
except ImportError:
    install('reportlab')
    from reportlab.pdfgen import canvas

text_content = """Tesco Summer Sale 2026!
Join us for our biggest event of the year. 
Get up to 50% off on fresh produce, clothing, and electronics.
Dates: June 1st - June 15th
Location: All Tesco Stores UK and online.
Special offer: Use code SUMMER26 for an extra 10% off your online cart.
We are bringing you the freshest deals!
"""

# Create DOCX
doc = Document()
doc.add_heading('Tesco Promotional Material', 0)
doc.add_paragraph(text_content)
doc.save('sample_promo.docx')
print('Created sample_promo.docx')

# Create PDF
c = canvas.Canvas('sample_promo.pdf')
c.drawString(100, 800, 'Tesco Promotional Material')
y = 760
for line in text_content.split('\n'):
    c.drawString(100, y, line)
    y -= 20
c.save()
print('Created sample_promo.pdf')
